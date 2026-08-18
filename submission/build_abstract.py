"""
Builds submission/HQIEA_ARGC_Abstract.docx: a 2-page conference/workshop
abstract following the format in abstract.txt (Title and authors / Background
and objectives / Methodology / Key results / Significance and impact /
3-5 keywords / References). Acknowledgements is omitted for now, per the
author's instruction -- add it back (a heading + text paragraph, same pattern
as the other sections) once there is something to acknowledge.

Content is grounded in logs.txt sections 14-15 (the full-scale 30-run/
n-gen=500 campaign and the plateau-diagnosis/restart-fix results) and
results/*_indicators.csv (the actual campaign data -- see generate_figure.py).
References are the five reference PDFs in the parent Paper1/ directory plus
the companion single-objective QIEA paper (same author group). Author list/
affiliation confirmed correct by the author (carried over from the companion
Paper 1 submission).

Not part of the research pipeline (see src/) -- a one-off document-generation
script for this deliverable. Re-run after editing the text below, or after
regenerating the figure, to rebuild the .docx.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

TITLE = (
    "A Hybrid Quantum-Inspired Evolutionary Algorithm with Adaptive "
    "Rotation-Gate Control for Many-Objective Vehicle Routing: Diagnosis, "
    "Correction, and Comparative Evaluation"
)

AUTHORS = "Olivian-Antonio Gîlea, Adrian Florea, Claudia Banciu, Alexandru Telcean"
AFFILIATION = (
    "Computer Science and Electrical Engineering Department, "
    '"Lucian Blaga" University of Sibiu, Sibiu 550025, Romania'
)

BACKGROUND = (
    "Quantum-inspired evolutionary algorithms (QIEAs) combine qubit-encoded "
    "chromosomes with rotation-gate operators to balance exploration and "
    "exploitation, and both quantum-inspired and true-quantum solvers have "
    "recently been applied to multi-objective optimization (Kotil et al., "
    "2025; King, 2025). However, these studies remain confined to synthetic "
    "benchmarks or low-dimensional (2-3 objective) problems, and a prior "
    "single-objective QIEA applied to real waste-collection routing found "
    "that its repair step for infeasible tours collapses population "
    "diversity as instance size grows (Gîlea et al., 2026). This work "
    "extends qubit-encoded optimization to a genuinely many-objective "
    "(5-objective), real-world setting — capacitated vehicle routing "
    "for urban waste collection — and asks whether a "
    "decomposition-guided QIEA with an adaptive rotation-gate control "
    "(ARGC) mechanism can close the gap to established many-objective "
    "evolutionary algorithms."
)

METHODOLOGY = (
    "We propose HQIEA-ARGC: a qubit-angle chromosome decoded either by "
    "argsort (permutation, for routing) or a sine-squared mapping "
    "(continuous, for synthetic benchmarks), eliminating the repair step "
    "that caused diversity collapse in the prior study. Many-objective "
    "scalability comes from MOEA/D-style Tchebycheff decomposition "
    "(Das-Dennis weight vectors, neighborhood-restricted mating); the "
    "rotation-gate step size decays across generations but is boosted, "
    "together with the mutation rate, whenever population diversity stalls "
    "— the ARGC mechanism. We formulate Sibiu’s waste-collection "
    "routing as a 5-objective CVRP (distance, time, cost, emissions, "
    "workload balance) and validate on the ZDT/DTLZ/WFG synthetic suites, "
    "four CVRPLIB instances, and three real Sibiu routes, against NSGA-II, "
    "SPEA2, MOEA/D, and RVEA (pymoo). We report the full-scale protocol: 30 "
    "independent runs and 500 generations per instance, scored by "
    "hypervolume, spacing and spread, with Wilcoxon and Friedman tests."
)

KEY_RESULTS = (
    "At full scale, HQIEA-ARGC initially lost to all four baselines on "
    "every one of the seven CVRP instances (hypervolume ratio to the best "
    "baseline: 0.17–0.45). A generation-wise diagnostic traced this to "
    "a hard, early hypervolume plateau: the ARGC stagnation-escape "
    "essentially never triggered under its original threshold, leaving the "
    "fixed-size decomposition population with no source of renewed "
    "diversity once converged. A plateau-gated diversity-reinjection "
    "scheme — reseeding half of the decomposition subproblems once the "
    "archive stalls for 10 generations — produced large, statistically "
    "significant hypervolume gains on all seven instances (+27.7% to "
    "+66.0%, Wilcoxon p<0.005 throughout), the most consistent effect found "
    "in our tuning study. Figure 1 shows the result: HQIEA-ARGC’s mean "
    "ratio to the best baseline rose to 0.48 (range 0.30–0.58), and it "
    "now statistically ties NSGA-II on one instance (route1_334, p=0.237) "
    "— its first tie against any baseline at full scale — while "
    "remaining significantly behind MOEA/D and RVEA specifically on all "
    "seven."
)

FIGURE_CAPTION = (
    "Figure 1. Mean hypervolume ratio of HQIEA-ARGC to the best-performing "
    "classical baseline, per instance, at the full validation protocol (30 "
    "runs, 500 generations). The highlighted bar (route1_334) marks the "
    "one instance where HQIEA-ARGC is statistically indistinguishable from "
    "a baseline (NSGA-II, p=0.237), even though MOEA/D still leads there "
    "in mean hypervolume."
)

SIGNIFICANCE = (
    "To our knowledge, this is the first evaluation of a qubit-encoded "
    "evolutionary algorithm on a genuinely many-objective (5-objective), "
    "real-world routing problem, rather than a synthetic or 2–3 "
    "objective case. Two findings generalize beyond this application: a "
    "repair-free permutation decode removes the diversity-collapse failure "
    "mode reported for qubit chromosomes at scale, and an explicit "
    "diagnose-then-fix methodology for a stalled stagnation-escape "
    "mechanism yields the largest, most reproducible gain in the study. "
    "The remaining, still-significant gap to decomposition- and "
    "reference-vector-guided baselines identifies population/decomposition "
    "granularity as the priority lever for future work bridging "
    "quantum-inspired representations with many-objective decomposition "
    "methods."
)

KEYWORDS = (
    "quantum-inspired evolutionary algorithm; many-objective optimization; "
    "capacitated vehicle routing; MOEA/D decomposition; adaptive "
    "rotation-gate control"
)

REFERENCES = [
    "Kotil, A., Pelofske, E., Riedmüller, S., Egger, D. J., Eidenbenz, "
    "S., Koch, T., & Woerner, S. (2025). Quantum Approximate "
    "Multi-Objective Optimization. arXiv:2503.22797 [quant-ph].",
    "King, A. D. (2025). Multi-objective optimization by quantum "
    "annealing. arXiv:2511.01762 [quant-ph].",
    "Koch, T., Bernal Neira, D. E., Chen, Y., Cortiana, G., Egger, D. J., "
    "Heese, R., et al. (2026). The Quantum Optimization Benchmarking "
    "Library. Nature Computational Science. "
    "https://doi.org/10.1038/s43588-026-00991-1",
    "Gîlea, O.-A., Florea, A., Banciu, C., & Telcean, A. (2026). "
    "Scalable Quantum-Inspired Evolutionary Algorithm for Multi-Objective "
    "Optimization: Real-World Problems. Manuscript in preparation, "
    '"Lucian Blaga" University of Sibiu. [confirm venue/status before citing]',
    "Ghlib, R., Bouhadouza, R., & Hnaien, F. (2026). Scalable "
    "multi-objective genetic algorithm for quantum circuit optimization. "
    "Scientific Reports, 16, 17977. "
    "https://doi.org/10.1038/s41598-026-47674-5",
]


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Cm(2.0))


def add_section_heading(doc, run_in_text, body_text):
    p = doc.add_paragraph()
    r = p.add_run(run_in_text + " ")
    r.bold = True
    p.add_run(body_text)


def add_figure(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(
        "submission/figures/hv_ratio_comparison.png", width=Cm(15.0)
    )
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(FIGURE_CAPTION)
    run.italic = True
    run.font.size = Pt(9)
    cap.paragraph_format.space_after = Pt(8)


def build():
    doc = Document()
    set_base_style(doc)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(13)

    authors_p = doc.add_paragraph()
    authors_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_p.paragraph_format.space_after = Pt(2)
    run = authors_p.add_run(AUTHORS)
    run.font.size = Pt(11)

    aff_p = doc.add_paragraph()
    aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff_p.paragraph_format.space_after = Pt(4)
    run = aff_p.add_run(AFFILIATION)
    run.italic = True
    run.font.size = Pt(9.5)

    add_section_heading(doc, "Background and objectives.", BACKGROUND)
    add_section_heading(doc, "Methodology.", METHODOLOGY)
    add_section_heading(doc, "Key results.", KEY_RESULTS)
    add_figure(doc)
    add_section_heading(doc, "Significance and impact.", SIGNIFICANCE)

    kw_p = doc.add_paragraph()
    kw_p.paragraph_format.space_after = Pt(6)
    r = kw_p.add_run("Keywords: ")
    r.bold = True
    kw_p.add_run(KEYWORDS)

    ref_heading = doc.add_paragraph()
    ref_heading.paragraph_format.space_after = Pt(3)
    r = ref_heading.add_run("References")
    r.bold = True

    for i, ref in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.add_run(f"[{i}] {ref}")
        for run in p.runs:
            run.font.size = Pt(9.5)

    doc.save("submission/HQIEA_ARGC_Abstract.docx")
    print("saved submission/HQIEA_ARGC_Abstract.docx")


if __name__ == "__main__":
    build()
